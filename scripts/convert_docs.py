"""
convert_docs.py
----------------
سكريبت لتحويل ملفات DOCX الموجودة في knowledge_base إلى ملفات نصية (TXT)
لتسهيل المعالجة والفهرسة لاحقاً.

المتطلبات:
    pip install python-docx

الاستخدام:
    python scripts/convert_docs.py
    python scripts/convert_docs.py --output data/cache/converted_texts
"""

import argparse
import logging
import sys
from pathlib import Path

try:
    import docx
except ImportError:
    print("مكتبة python-docx غير مثبتة. رجاء تشغيل: pip install python-docx")
    sys.exit(1)

# ============================================
# المسارات الأساسية
# ============================================
ROOT_DIR = Path(__file__).resolve().parent.parent
KNOWLEDGE_BASE_DIR = ROOT_DIR / "knowledge_base"
DEFAULT_OUTPUT_DIR = ROOT_DIR / "data" / "cache" / "converted_texts"

KNOWLEDGE_SUBDIRS = ["policies", "contracts", "quotations", "quality_reports"]

# ============================================
# إعداد نظام التسجيل (Logging)
# ============================================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("convert_docs")


def extract_text_from_docx(file_path: Path) -> str:
    """يستخرج النص الكامل من ملف DOCX (فقرات + جداول)."""
    document = docx.Document(str(file_path))
    parts = []

    # استخراج الفقرات
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    # استخراج الجداول
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def convert_all(output_dir: Path, overwrite: bool) -> dict:
    """يحول كل ملفات DOCX الموجودة في knowledge_base إلى ملفات TXT."""
    stats = {"converted": 0, "skipped": 0, "failed": 0}

    for subdir in KNOWLEDGE_SUBDIRS:
        source_dir = KNOWLEDGE_BASE_DIR / subdir
        if not source_dir.exists():
            logger.warning("المجلد غير موجود، تم تجاوزه: %s", source_dir)
            continue

        target_dir = output_dir / subdir
        target_dir.mkdir(parents=True, exist_ok=True)

        docx_files = sorted(source_dir.glob("*.docx"))
        logger.info("جاري معالجة %d ملف في %s", len(docx_files), subdir)

        for file_path in docx_files:
            output_path = target_dir / f"{file_path.stem}.txt"

            if output_path.exists() and not overwrite:
                logger.info("تم تجاوز (موجود بالفعل): %s", output_path.name)
                stats["skipped"] += 1
                continue

            try:
                text = extract_text_from_docx(file_path)
                output_path.write_text(text, encoding="utf-8")
                logger.info("تم التحويل: %s → %s", file_path.name, output_path.relative_to(ROOT_DIR))
                stats["converted"] += 1
            except Exception as exc:
                logger.error("فشل تحويل %s: %s", file_path.name, exc)
                stats["failed"] += 1

    return stats


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="تحويل ملفات DOCX إلى TXT")
    parser.add_argument(
        "--output",
        type=str,
        default=None,
        help="مسار مجلد الإخراج (افتراضياً: data/cache/converted_texts)",
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="إعادة تحويل الملفات حتى لو كانت نسخة TXT موجودة بالفعل",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output_dir = Path(args.output) if args.output else DEFAULT_OUTPUT_DIR
    output_dir.mkdir(parents=True, exist_ok=True)

    logger.info("بدء تحويل ملفات DOCX إلى TXT...")
    logger.info("مجلد الإخراج: %s", output_dir)

    stats = convert_all(output_dir, args.overwrite)

    logger.info(
        "انتهت العملية ✅ — تم التحويل: %d | تم التجاوز: %d | فشل: %d",
        stats["converted"],
        stats["skipped"],
        stats["failed"],
    )


if __name__ == "__main__":
    main()
